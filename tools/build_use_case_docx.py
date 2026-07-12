from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "use-case-full-flow-zrinttailor-desktop-dengan-penjelasan.docx"
SHOT = ROOT / "docs" / "screenshots"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
LIGHT_FILL = "F2F4F7"
SOFT_BLUE = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="111827")
        set_cell_shading(hdr[i], SOFT_BLUE)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_image(doc, filename, caption, width=6.1, explanation=None):
    path = SHOT / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    if explanation:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(4)
        p_title.paragraph_format.space_after = Pt(3)
        r_title = p_title.add_run("Penjelasan " + caption.split(".")[0])
        r_title.bold = True
        r_title.font.size = Pt(10)
        r_title.font.color.rgb = RGBColor.from_string(DARK_BLUE)

        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_after = Pt(8)
        r_body = p_body.add_run(explanation)
        r_body.font.size = Pt(10)
        r_body.font.color.rgb = RGBColor.from_string(INK)


def set_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("DOKUMENTASI USE CASE DAN FULL FLOW")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ZRINTTAILOR")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Pemisahan Alur User/Pelanggan dan Admin untuk Sistem Intelligent Tailoring Service"
    )
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    add_note(
        doc,
        "Konteks dokumen",
        "Dokumen ini menjelaskan use case dan tutorial full flow aplikasi ZRINTTAILOR. "
        "Bagian User/Pelanggan berfokus pada ukur badan, pemesanan, pembayaran, tracking, chat, dan rating. "
        "Bagian Admin berfokus pada pengelolaan layanan, katalog, pesanan, pembayaran, pengiriman, chat, dan konten."
    )

    add_table(
        doc,
        ["Metadata", "Keterangan"],
        [
            ["Tanggal", "12 Juli 2026"],
            ["Lingkungan", "Laravel lokal http://localhost:8001"],
            ["Akun user demo", "demo@zrinttailor.com / password123"],
            ["Akun admin demo", "admin@zrinttailor.com / password123"],
            ["Jenis dokumen", "Use case, tutorial alur, dan dokumentasi penjelasan"],
        ],
        widths=[1.8, 4.5],
    )
    doc.add_page_break()


def build_doc():
    doc = Document()
    set_styles(doc)
    add_cover(doc)

    doc.add_heading("1. Penjelasan: Ini Use Case User atau Admin?", level=1)
    doc.add_paragraph(
        "Flow yang sebelumnya dibuat adalah flow User/Pelanggan, karena isinya menjelaskan cara pelanggan login, "
        "mengukur badan, dan membuat pesanan. Agar dokumentasi tugas akhir lebih lengkap, flow perlu dipisahkan "
        "menjadi dua aktor utama: User/Pelanggan dan Admin."
    )
    add_table(
        doc,
        ["Aktor", "Peran Utama", "Contoh Fitur"],
        [
            [
                "User/Pelanggan",
                "Menggunakan layanan jahit dari sisi pemesan.",
                "Registrasi/login, ukur badan CV, buat pesanan, bayar, tracking, chat admin, rating.",
            ],
            [
                "Admin",
                "Mengelola operasional layanan jahit dan proses pesanan.",
                "Dashboard, layanan, katalog, pesanan, pembayaran, upload desain/resi, chat, settings.",
            ],
        ],
        widths=[1.4, 2.4, 2.4],
    )

    doc.add_heading("2. Use Case User/Pelanggan", level=1)
    add_table(
        doc,
        ["Kode", "Use Case", "Tujuan", "Hasil Akhir"],
        [
            ["UC-U01", "Login/Registrasi", "Masuk ke sistem sebagai pelanggan.", "User masuk ke dashboard pelanggan."],
            ["UC-U02", "Ukur Badan CV", "Upload foto full body dan benda referensi.", "Sistem menganalisis ukuran dalam cm."],
            ["UC-U03", "Buat Pesanan", "Memesan layanan jahit custom/permak/desain.", "Pesanan masuk ke sistem admin."],
            ["UC-U04", "Pembayaran", "Upload bukti pembayaran.", "Pembayaran menunggu verifikasi admin."],
            ["UC-U05", "Tracking Pesanan", "Melihat status produksi dan pengiriman.", "User mengetahui progres pesanan."],
            ["UC-U06", "Chat Admin", "Konsultasi kebutuhan dan status pesanan.", "Komunikasi user-admin tercatat."],
            ["UC-U07", "Rating/Testimoni", "Memberi ulasan setelah pesanan selesai.", "Testimoni masuk ke sistem."],
        ],
        widths=[0.7, 1.5, 2.2, 2.0],
    )

    doc.add_heading("2.1 Tutorial Flow User/Pelanggan", level=2)
    add_numbers(
        doc,
        [
            "Buka halaman login dan masuk memakai akun pelanggan.",
            "Masuk ke dashboard pelanggan.",
            "Buka menu Ukur Badan (CV).",
            "Upload satu foto full body dan pilih benda referensi.",
            "Kirim form ke backend untuk analisis ukuran.",
            "Buka halaman Buat Pesanan.",
            "Pilih layanan Jahit Custom jika mengikuti flow utama proposal.",
            "Pilih jenis pakaian sesuai proposal: Kemeja, Baju Dinas, Baju Sekolah, Baju Koko, Kebaya, Gamis, Celana Kain, atau Rok Kain.",
            "Pilih ukuran dari hasil CV atau isi manual dalam sentimeter.",
            "Lengkapi alamat dan nomor penerima.",
            "Klik Buat Pesanan.",
            "Lanjutkan pembayaran dan tracking setelah pesanan dikonfirmasi admin.",
        ],
    )
    add_note(
        doc,
        "Catatan proposal",
        "Untuk layanan custom, sistem tidak memakai ukuran standar S/M/L/XL sebagai dasar utama. Ukuran yang digunakan adalah ukuran tubuh dalam sentimeter."
    )

    add_image(
        doc,
        "01-login.png",
        "Gambar 1. Halaman login user.",
        width=5.8,
        explanation=(
            "Gambar ini menunjukkan pintu masuk aplikasi untuk aktor User/Pelanggan. "
            "Pada sisi kanan terdapat form autentikasi berisi alamat email, password, opsi ingat saya, "
            "tautan lupa password, tombol masuk, dan opsi masuk dengan Google. Sisi kiri menampilkan branding "
            "ZRINTTAILOR serta ringkasan nilai sistem: pengukuran otomatis melalui foto, pemantauan progres pesanan, "
            "dan notifikasi update status. Tampilan ini membuktikan bahwa pelanggan harus login terlebih dahulu "
            "sebelum mengakses fitur pemesanan dan ukur badan."
        ),
    )
    add_image(
        doc,
        "02-dashboard-pesan.png",
        "Gambar 2. Dashboard/area pelanggan setelah login.",
        width=5.8,
        explanation=(
            "Gambar ini memperlihatkan area pelanggan setelah login berhasil. Sidebar kiri berisi navigasi utama "
            "untuk Pesanan Saya, Chat Admin, dan Ukur Badan (CV). Area konten menampilkan halaman Buat Pesanan "
            "dengan langkah awal pemilihan layanan, seperti Permak, Desain Digital, dan Jahit Custom. Pada flow proposal, "
            "layanan Jahit Custom menjadi alur paling relevan karena berhubungan langsung dengan rekomendasi ukuran "
            "dan pemesanan pakaian custom."
        ),
    )
    add_image(
        doc,
        "03-ukur-badan.png",
        "Gambar 3. Halaman ukur badan berbasis Computer Vision.",
        width=5.8,
        explanation=(
            "Gambar ini menunjukkan fitur Ukur Badan (AI/CV) dari sisi pelanggan. Bagian atas berisi panduan foto: "
            "satu orang berdiri tegak, seluruh tubuh terlihat, benda referensi terlihat jelas, pencahayaan cukup, "
            "dan foto tidak buram. Di bagian form, pelanggan memilih benda referensi seperti kertas A4, kartu ATM/KTP, "
            "atau benda custom dengan ukuran diketahui. Form ini juga menyediakan upload foto full body dan mengirim data "
            "ke backend untuk proses analisis ukuran, sehingga perhitungan utama tidak lagi dilakukan sebagai logika "
            "MediaPipe penuh di browser."
        ),
    )
    add_image(
        doc,
        "04-buat-pesanan.png",
        "Gambar 4. Halaman buat pesanan user.",
        width=5.8,
        explanation=(
            "Gambar ini menampilkan halaman pembuatan pesanan dari sisi User/Pelanggan. Setelah memilih layanan, "
            "pelanggan mengisi detail pakaian, termasuk jenis pakaian, warna, bahan, deskripsi kebutuhan, dan gambar referensi. "
            "Jenis pakaian yang disediakan mengikuti ruang lingkup proposal, yaitu Kemeja, Baju Dinas, Baju Sekolah, "
            "Baju Koko, Kebaya, Gamis, Celana Kain, dan Rok Kain. Opsi denim/jeans tidak ditampilkan agar tetap sesuai "
            "dengan batasan proposal."
        ),
    )
    add_image(
        doc,
        "05-input-ukuran-manual.png",
        "Gambar 5. Input ukuran manual dalam sentimeter.",
        width=5.8,
        explanation=(
            "Gambar ini memperlihatkan bagian pemilihan sumber ukuran pada proses pemesanan. Pelanggan dapat memakai "
            "hasil ukur dari foto Computer Vision atau mengisi ukuran manual. Field manual menggunakan satuan sentimeter, "
            "seperti dada, pinggang, pinggul, lebar bahu, panjang lengan, dan tinggi badan. Bagian ini penting karena "
            "menegaskan bahwa layanan custom tidak memakai ukuran standar S/M/L/XL sebagai dasar produksi, melainkan "
            "ukuran tubuh aktual pelanggan."
        ),
    )

    doc.add_page_break()
    doc.add_heading("3. Use Case Admin", level=1)
    add_table(
        doc,
        ["Kode", "Use Case", "Tujuan", "Hasil Akhir"],
        [
            ["UC-A01", "Dashboard Admin", "Melihat ringkasan bisnis dan status pesanan.", "Admin memahami kondisi operasional."],
            ["UC-A02", "Kelola Layanan", "Menambah/mengubah layanan seperti custom, permak, desain.", "Layanan tampil ke user."],
            ["UC-A03", "Kelola Katalog", "Mengelola contoh produk/model jahitan.", "Katalog dapat dipilih user."],
            ["UC-A04", "Kelola Pesanan", "Melihat detail pesanan, ukuran, alamat, dan status.", "Pesanan dapat diproses."],
            ["UC-A05", "Konfirmasi Harga dan Status", "Menentukan harga dan memperbarui progres.", "User mendapat status terbaru."],
            ["UC-A06", "Verifikasi Pembayaran", "Memeriksa bukti bayar user.", "Pembayaran diterima atau ditolak."],
            ["UC-A07", "Upload Desain/Resi", "Mengirim file desain atau nomor resi.", "User dapat mengunduh file atau tracking kiriman."],
            ["UC-A08", "Chat User", "Melayani pertanyaan pelanggan.", "Komunikasi operasional berjalan."],
            ["UC-A09", "Kelola Testimoni dan Settings", "Mengatur ulasan, QR/DANA, dan konfigurasi.", "Informasi bisnis tetap akurat."],
        ],
        widths=[0.7, 1.6, 2.1, 2.0],
    )

    doc.add_heading("3.1 Tutorial Flow Admin", level=2)
    add_numbers(
        doc,
        [
            "Login sebagai admin.",
            "Buka Dashboard untuk melihat total pesanan, pendapatan, total pelanggan, dan pembayaran menunggu verifikasi.",
            "Kelola data layanan agar pilihan layanan user tetap sesuai kebutuhan bisnis.",
            "Kelola katalog sebagai referensi model pakaian.",
            "Buka menu Pesanan untuk melihat pesanan masuk dari pelanggan.",
            "Masuk ke detail pesanan untuk membaca data pelanggan, detail pakaian, ukuran badan, alamat, dan status pembayaran.",
            "Jika pesanan masih pending, admin mengonfirmasi pesanan dan menentukan harga.",
            "Jika user mengupload bukti pembayaran, admin memverifikasi atau menolak pembayaran.",
            "Admin memperbarui status pengerjaan: dikonfirmasi, diproses, selesai dibuat, dikirim, atau selesai.",
            "Untuk layanan desain, admin mengupload file desain.",
            "Untuk layanan fisik, admin menginput nomor resi pengiriman.",
            "Admin menjawab chat user jika ada konsultasi atau kendala.",
        ],
    )
    add_image(
        doc,
        "06-admin-dashboard.png",
        "Gambar 6. Dashboard admin.",
        width=5.8,
        explanation=(
            "Gambar ini menunjukkan dashboard untuk aktor Admin. Dashboard menampilkan ringkasan operasional seperti "
            "total pesanan, pendapatan bulan ini, total pelanggan, pembayaran yang menunggu verifikasi, distribusi status "
            "pesanan, grafik pesanan per bulan, grafik pendapatan, pesanan terbaru, dan pembayaran menunggu. Tampilan ini "
            "digunakan admin untuk memantau kondisi bisnis dan menentukan tindakan operasional berikutnya."
        ),
    )
    add_image(
        doc,
        "07-admin-orders.png",
        "Gambar 7. Daftar pesanan admin.",
        width=5.8,
        explanation=(
            "Gambar ini memperlihatkan menu Pesanan pada sisi Admin. Admin dapat melihat daftar pesanan masuk, mencari "
            "pesanan, memfilter status, melihat kode pesanan, nama pelanggan, jenis layanan, tanggal, status, dan total harga. "
            "Use case ini menjadi titik awal admin untuk membuka detail pesanan, mengecek ukuran pelanggan, mengonfirmasi harga, "
            "memproses pengerjaan, dan memperbarui status pesanan."
        ),
    )
    add_image(
        doc,
        "08-admin-services.png",
        "Gambar 8. Kelola layanan admin.",
        width=5.8,
        explanation=(
            "Gambar ini menampilkan halaman Kelola Layanan dari sisi Admin. Admin dapat menambah, mengubah, mengaktifkan, "
            "menonaktifkan, atau menghapus layanan seperti Permak, Desain Digital, dan Jahit Custom. Data layanan ini "
            "berpengaruh langsung pada pilihan yang muncul di halaman buat pesanan User/Pelanggan, sehingga admin berperan "
            "menjaga agar layanan yang tersedia tetap sesuai kebutuhan bisnis dan ruang lingkup sistem."
        ),
    )

    doc.add_heading("4. Hubungan Flow User dan Admin", level=1)
    add_table(
        doc,
        ["Tahap", "Aksi User/Pelanggan", "Aksi Admin", "Output"],
        [
            ["1", "Login dan buat pesanan.", "Menerima pesanan di menu admin.", "Pesanan masuk."],
            ["2", "Upload foto ukur badan atau isi ukuran manual.", "Melihat ukuran pada detail pesanan.", "Data ukuran tersedia."],
            ["3", "Menunggu konfirmasi.", "Konfirmasi pesanan dan isi harga.", "Tagihan siap dibayar."],
            ["4", "Upload bukti pembayaran.", "Verifikasi/tolak pembayaran.", "Status pembayaran diperbarui."],
            ["5", "Menunggu proses jahit/desain.", "Update status pengerjaan.", "Progress terlihat di user."],
            ["6", "Menerima hasil/produk.", "Upload file desain atau input resi.", "User mendapat file/resi."],
            ["7", "Konfirmasi selesai dan beri rating.", "Melihat testimoni.", "Pesanan selesai."],
        ],
        widths=[0.5, 1.9, 1.9, 1.7],
    )

    doc.add_heading("5. Kesesuaian dengan Proposal", level=1)
    add_bullets(
        doc,
        [
            "Fitur ukur badan menggunakan foto full body dan benda referensi.",
            "Analisis ukuran diarahkan ke backend melalui endpoint analisis.",
            "Ukuran disimpan dan digunakan dalam satuan sentimeter.",
            "Pakaian yang dipilih user mengikuti ruang lingkup proposal.",
            "Celana dan rok diarahkan sebagai kain, bukan denim/jeans.",
            "Preset ukuran standar S/M/L/XL tidak digunakan sebagai dasar custom tailoring.",
        ],
    )

    doc.add_heading("6. Rekomendasi Tambahan untuk Tugas Akhir", level=1)
    add_note(
        doc,
        "Modul evaluasi akurasi",
        "Agar dokumen tugas akhir lebih kuat, tambahkan data ground truth dari ukuran manual penjahit dan bandingkan dengan hasil Computer Vision. Nilai MAE dapat digunakan sebagai metrik evaluasi."
    )
    add_table(
        doc,
        ["Foto", "Atribut", "Hasil CV", "Ground Truth", "Error Absolut"],
        [
            ["sample-001", "Dada", "92.0", "91.0", "1.0"],
            ["sample-001", "Pinggang", "78.0", "80.0", "2.0"],
        ],
        widths=[1.2, 1.4, 1.1, 1.4, 1.3],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("ZRINTTAILOR - Dokumentasi Use Case dan Full Flow")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
